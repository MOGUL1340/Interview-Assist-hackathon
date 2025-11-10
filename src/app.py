from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import base64
import logging
import io
from resume_analyzer import process_resume
from audio_transcriber import process_meeting_recording, extract_meeting_insights
from interview_plan_generator import create_complete_interview_plan
from code_challenge_generator import create_challenge_suite
from excel_generator import create_interview_excel
from dotenv import load_dotenv

# Document parsing imports
try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    import docx2txt
    DOC_SUPPORT = True
except ImportError:
    DOC_SUPPORT = False

app = Flask(__name__)
CORS(app)

# Load .env from project root (parent directory)
load_dotenv(dotenv_path=os.path.join(os.path.dirname(__file__), '..', '.env'))

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def extract_text_from_file(file_bytes, filename):
    """
    Extract text from various file formats (TXT, PDF, DOC, DOCX).

    Args:
        file_bytes: Raw bytes of the file
        filename: Name of the file (used to determine format)

    Returns:
        Extracted text as string, or None if extraction failed
    """
    file_extension = filename.lower().split('.')[-1]

    # TXT files - try multiple encodings
    if file_extension in ['txt', 'text']:
        for encoding in ['utf-8', 'utf-8-sig', 'latin-1', 'windows-1252', 'cp1252']:
            try:
                text = file_bytes.decode(encoding)
                logging.info(f"Successfully decoded TXT file using {encoding} encoding")
                return text
            except (UnicodeDecodeError, AttributeError):
                continue
        logging.error("Failed to decode TXT file with any encoding")
        return None

    # PDF files
    elif file_extension == 'pdf':
        if not PDF_SUPPORT:
            logging.error("PDF support not available. Install PyPDF2.")
            return None
        try:
            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            logging.info(f"Successfully extracted text from PDF ({len(pdf_reader.pages)} pages)")
            return text.strip()
        except Exception as e:
            logging.error(f"Error extracting text from PDF: {str(e)}")
            return None

    # DOCX files
    elif file_extension == 'docx':
        if not DOCX_SUPPORT:
            logging.error("DOCX support not available. Install python-docx.")
            return None
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)

            text = "\n".join(text_parts)
            logging.info(f"Successfully extracted text from DOCX ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")
            logging.info(f"Extracted text preview (first 200 chars): {text[:200]}")
            return text.strip()
        except Exception as e:
            logging.error(f"Error extracting text from DOCX: {str(e)}")
            return None

    # DOC files (older Word format)
    elif file_extension == 'doc':
        if not DOC_SUPPORT:
            logging.error("DOC support not available. Install docx2txt.")
            return None
        try:
            # Save to temp file since docx2txt needs a file path
            temp_path = os.path.join(UPLOAD_FOLDER, f"temp_{os.urandom(8).hex()}.doc")
            with open(temp_path, 'wb') as f:
                f.write(file_bytes)
            text = docx2txt.process(temp_path)
            os.remove(temp_path)
            logging.info(f"Successfully extracted text from DOC file")
            return text.strip()
        except Exception as e:
            logging.error(f"Error extracting text from DOC: {str(e)}")
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return None

    else:
        logging.error(f"Unsupported file format: {file_extension}")
        return None

@app.route('/generate_interview_plan', methods=['POST'])
def generate_interview_plan_endpoint():
    """
    Main endpoint for generating interview preparation plan.

    Expected JSON payload:
    {
        "candidate_cv": {
            "name": "resume.pdf",
            "content": "base64_encoded_content"
        },
        "meeting_recording": {
            "name": "meeting.mp3",  // optional if transcript provided
            "content": "base64_encoded_content"
        },
        "meeting_transcript": "text transcript...",  // alternative to recording
        "job_requirements": "...",
        "job_position": "Senior .NET Developer",
        "interview_duration_minutes": 30
    }
    """
    try:
        logging.info("Received interview plan generation request")
        data = request.json

        # 1. Process Resume
        logging.info("Step 1: Processing resume")
        resume = data.get('candidate_cv', {})
        if not resume.get('content'):
            return jsonify({'status': 'error', 'message': 'Resume is required'}), 400

        # Extract text from resume file (supports TXT, PDF, DOC, DOCX)
        resume_bytes = base64.b64decode(resume['content'])
        resume_filename = resume.get('name', 'resume.txt')

        resume_text = extract_text_from_file(resume_bytes, resume_filename)

        if not resume_text or len(resume_text.strip()) < 10:
            logging.error(f"Resume text is too short or empty. Length: {len(resume_text) if resume_text else 0}")
            return jsonify({'status': 'error', 'message': 'Resume file appears to be empty or contains insufficient text. Please provide a resume with at least some content.'}), 400

        logging.info(f"Resume text extracted successfully. Length: {len(resume_text)} characters")
        resume_analysis = process_resume(resume_text)

        if "error" in resume_analysis:
            return jsonify({'status': 'error', 'message': resume_analysis['error']}), 500

        # 2. Process Meeting Recording or Transcript
        logging.info("Step 2: Processing meeting information")
        meeting_insights = {}

        if data.get('meeting_transcript'):
            # Use provided transcript
            transcript_text = data.get('meeting_transcript')
            meeting_insights = extract_meeting_insights(transcript_text)
        elif data.get('meeting_recording'):
            # Transcribe recording
            recording = data.get('meeting_recording', {})
            file_extension = recording.get('name', 'audio.mp3').split('.')[-1]

            meeting_data = process_meeting_recording(
                recording.get('content'),
                is_base64=True,
                file_extension=file_extension
            )

            if "error" in meeting_data:
                return jsonify({'status': 'error', 'message': meeting_data['error']}), 500

            meeting_insights = meeting_data.get('insights', {})
        else:
            # No meeting data, use defaults
            meeting_insights = {
                "job_requirements": data.get('job_requirements', ''),
                "interview_duration_minutes": data.get('interview_duration_minutes', 30),
                "topics_to_cover": [],
                "code_challenge_needed": True
            }

        # 3. Get Job Details
        logging.info("Step 3: Processing job details")
        job_position = data.get('job_position', '')

        job_details = {
            "title": job_position or "Position",
            "description": data.get('job_requirements', 'No specific requirements provided')
        }

        # 4. Generate Interview Plan
        logging.info("Step 4: Generating interview plan")
        interview_duration = data.get('interview_duration_minutes',
                                     meeting_insights.get('interview_duration_minutes', 30))

        interview_plan = create_complete_interview_plan(
            resume_analysis,
            {'insights': meeting_insights},
            job_details,
            interview_duration
        )

        if "error" in interview_plan:
            return jsonify({'status': 'error', 'message': interview_plan['error']}), 500

        # 5. Generate Code Challenges (only if requested)
        logging.info("Step 5: Generating code challenges")
        include_code_challenges = data.get('include_code_challenges', False)
        if include_code_challenges:
            code_challenges = create_challenge_suite(
                job_details,
                resume_analysis,
                {'insights': meeting_insights}
            )
        else:
            code_challenges = {"coding_challenges": [], "system_design": None, "debugging_challenge": None}
            logging.info("Code challenges skipped - not requested by user")

        # 6. Generate Excel File
        logging.info("Step 6: Generating Excel file")
        excel_path = create_interview_excel(interview_plan, code_challenges)

        if not excel_path:
            return jsonify({'status': 'error', 'message': 'Failed to generate Excel file'}), 500

        # Read Excel file as base64 for sending to frontend
        with open(excel_path, 'rb') as f:
            excel_content = base64.b64encode(f.read()).decode('utf-8')

        # Clean up Excel file
        try:
            os.remove(excel_path)
        except:
            pass

        response_data = {
            'status': 'success',
            'interview_plan': interview_plan,
            'code_challenges': code_challenges,
            'excel_file': {
                'name': os.path.basename(excel_path),
                'content': excel_content
            }
        }

        # Log the structure for debugging
        logging.info(f"Interview plan keys: {list(interview_plan.keys())}")
        logging.info(f"Metadata: {interview_plan.get('metadata', {})}")
        logging.info(f"Has prioritized_topics: {'prioritized_topics' in interview_plan}")
        logging.info(f"Code challenges keys: {list(code_challenges.keys()) if code_challenges else 'None'}")

        logging.info("Interview plan generated successfully")
        return jsonify(response_data), 200

    except Exception as e:
        logging.error(f"Error generating interview plan: {str(e)}")
        return jsonify({'status': 'error', 'message': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)